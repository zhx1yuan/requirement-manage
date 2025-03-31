# 入口文件
from fastapi import FastAPI, Depends, HTTPException, status, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.responses import JSONResponse
from sqlalchemy.exc import SQLAlchemyError

# import uvicorn

from typing import List, Optional
from fastapi.staticfiles import StaticFiles

from datetime import timedelta, datetime
from sqlalchemy.orm import Session
from sqlalchemy import or_, and_
from backend.database import engine, get_db

from backend import models, schemas, crud, auth, exceptions
from backend.models import PermissionLevel

from fastapi.responses import FileResponse
from docx import Document as DocxDocument
import os
from docx.shared import Pt
import html2text

###############     创建异步后端服务器实例，设置跨域资源共享      ##################

app = FastAPI()
app.mount('/swagger-ui-master', StaticFiles(directory='swagger-ui-master'))

# 允许的源列表
origins = [
    "http://localhost:5173",
    "http://localhost:5174",
    "http://localhost:3000",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:5174",
    "http://127.0.0.1:3000",
    # 可以添加更多允许的源
]

# 添加 CORS 中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

###############     创建数据库会话      ##################
# 根据预定义的表，创建所有表
models.Base.metadata.create_all(bind=engine)


##########################     API端点设置      ##########################

###############     用户注册、登录、获取token      #################

@app.post("/users/", response_model=schemas.User)
def create_user(user: schemas.UserCreate, db: Session = Depends(get_db)):
    # 先查询是否有已创建的user
    db_user = crud.get_user_by_username(db, username=user.username)
    # 有就返回400
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    # 没有就创建新的
    return crud.create_user(db=db, user=user)


# 登录
@app.post("/token", response_model=schemas.Token)
async def login_for_access_token(db: Session = Depends(get_db), form_data: OAuth2PasswordRequestForm = Depends()):
    user = auth.authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = auth.create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}


@app.get("/users/me/", response_model=schemas.User)
async def read_users_me(current_user: schemas.User = Depends(auth.get_current_active_user)):
    return current_user


###############     测试      ##################
# def get_current_user(db: Session = Depends(SessionLocal1):
#     return db

# def get_current_user_wrapper(token: str = Depends(oauth2_scheme)):
#     print("111111111111111111111111111111111111111111"+token)
#     payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
#     username: str = payload.get("sub")
#     print("111111111111111111111111111111111111111122" + username)
#     user = crud.get_user_by_username(db=Depends(get_db), username=username)
#
#     print("111111111111111111111111111111111111111122" + user)
#     return auth.get_current_user(token=token)
# @app.get("/test-current-user/", response_model=schemas.User)
# async def read_users_me(current_user: schemas.User = Depends(get_current_user_wrapper)):
#     return current_user

# @app.get("/test-current-user/")
# async def test_current_user(user=Depends(auth.get_current_user)):
#     return {"message": "Current user fetched successfully"}


###############     项目      ##################

@app.post("/projects/", response_model=schemas.Project)
def create_project(project: schemas.ProjectCreate, db: Session = Depends(get_db),
                   current_user: schemas.User = Depends(auth.get_current_active_user)):
    return crud.create_project(db=db, project=project, user_id=current_user.id)


@app.get("/projects/", response_model=List[schemas.Project])
def get_projects(db: Session = Depends(get_db), current_user: schemas.User = Depends(auth.get_current_active_user)):
    return crud.get_projects(db=db, user_id=current_user.id)




###############     列表      ##################

# 创建列表
@app.post("/lists/", response_model=schemas.ListSchema)
def create_list(project_id: int, list: schemas.ListCreate, db: Session = Depends(get_db)):
    return crud.create_list(db=db, list=list, project_id=project_id)

# 获取项目的完整列表树
@app.get("/projects/{project_id}/list-tree/", response_model=List[schemas.ListTree])
def get_project_list_tree(project_id: int, db: Session = Depends(get_db)):
    return crud.get_list_tree(db=db, project_id=project_id)

# 查询列表——根据列表名称、id、项目id
@app.get("/lists/", response_model=List[schemas.ListSchema])
def get_list(list_id: int = None, list_name: str = None, list_project_id: int = None, db: Session = Depends(get_db)):
    db_list = crud.get_a_list(db=db, list_id=list_id, list_name=list_name, list_project_id=list_project_id)
    if not db_list:
        raise HTTPException(status_code=404, detail="List not found")
    return db_list

# 列表更新
@app.put("/lists/{list_id}", response_model=schemas.ListSchema)
def update_list(
    list_id: int,
    list: schemas.ListCreate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """更新列表"""
    db_list = crud.update_list(db=db, list_id=list_id, list_data=list, user_id=current_user.id)
    if not db_list:
        raise HTTPException(status_code=404, detail="列表不存在或没有权限")
    return db_list

# 列表删除
@app.delete("/lists/{list_id}/")
def delete_list(list_id: int, db: Session = Depends(get_db)):
    deleted_list = crud.delete_list(db=db, list_id=list_id)
    if not deleted_list:
        raise HTTPException(status_code=404, detail="List not found")
    return {"message": "List deleted successfully"}


# 查询列表——根据项目id
# @app.get("/lists_by_project/", response_model=List[schemas.ListSchema])
# def get_list(list_project_id: int, db: Session = Depends(get_db)):
#     db_list = crud.get_all_list(db=db,list_project_id = list_project_id)
#     if not db_list:
#         raise HTTPException(status_code=404, detail="List not found")
#     return db_list




###############     文档      ##################

# 文档创建
@app.post("/lists/{list_id}/documents/", response_model=schemas.Document)
def create_document(
    list_id: int,
    document: schemas.DocumentCreate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    return crud.create_document(db=db, document=document, list_id=list_id, user_id=current_user.id)

# 文档查询
@app.get("/lists/{list_id}/documents/", response_model=List[schemas.Document])
def get_documents(
    list_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    return crud.get_documents(db=db, list_id=list_id)

# 获取单个文档
@app.get("/documents/{document_id}/", response_model=schemas.Document)
def get_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    document = crud.get_document(db=db, document_id=document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document

# 文档更新
@app.put("/documents/{document_id}/", response_model=schemas.Document)
def update_document(
    document_id: int,
    document: schemas.DocumentCreate,
    comment: str = None,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    updated_document = crud.update_document(
        db=db,
        document_id=document_id,
        document=document,
        user_id=current_user.id,
        comment=comment
    )
    if not updated_document:
        raise HTTPException(status_code=404, detail="Document not found")
    return updated_document

# 文档删除
@app.delete("/documents/{document_id}/")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    deleted_document = crud.delete_document(db=db, document_id=document_id)
    if not deleted_document:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted successfully"}

# 获取文档版本历史
@app.get("/documents/{document_id}/versions/", response_model=List[schemas.DocumentVersion])
def get_document_versions(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    return crud.get_document_versions(db=db, document_id=document_id)

# 恢复到指定版本
@app.post("/documents/{document_id}/restore/{version}/", response_model=schemas.Document)
def restore_document_version(
    document_id: int,
    version: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    restored_document = crud.restore_document_version(
        db=db,
        document_id=document_id,
        version=version,
        user_id=current_user.id
    )
    if not restored_document:
        raise HTTPException(status_code=404, detail="Document or version not found")
    return restored_document


###############     文档导出      ##################

# 创建HTML转换器
h = html2text.HTML2Text()
h.ignore_links = True
h.ignore_images = True
h.ignore_emphasis = True
h.body_width = 0  # 不限制行宽

def convert_html_to_text(html_content: str) -> str:
    """将HTML内容转换为纯文本"""
    if not html_content:
        return ""
    return h.handle(html_content).strip()

@app.post("/projects/{project_id}/export/")
def export_project(
    project_id: int,
    selected_list_ids: List[int],
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """导出项目中的文档"""
    # 验证项目是否存在
    project = db.query(models.Project).filter(
        models.Project.id == project_id,
        models.Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # 创建Word文档
    doc = DocxDocument()
    
    # 设置文档默认字体为宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 设置标题字体
    for i in range(1, 10):
        style = doc.styles[f'Heading {i}']
        style.font.name = '宋体'
        style.font.size = Pt(16 - i)
    
    # 添加项目标题
    title = doc.add_heading(f"项目：{project.name}", level=0)
    title.style.font.name = '宋体'
    title.style.font.size = Pt(20)
    
    # 添加导出信息
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run("导出时间：").bold = True
    info_paragraph.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    info_paragraph.add_run("\n导出用户：").bold = True
    info_paragraph.add_run(current_user.username)
    doc.add_paragraph()  # 添加空行

    # 获取所有选中的列表
    selected_lists = db.query(models.mylist).filter(
        models.mylist.id.in_(selected_list_ids),
        models.mylist.project_id == project_id
    ).all()

    # 按层级排序列表
    def sort_lists_by_level(lists):
        # 首先按层级排序
        sorted_lists = sorted(lists, key=lambda x: x.level_id)
        # 然后按同级顺序排序
        return sorted(sorted_lists, key=lambda x: (x.level_id, x.order_index))

    # 处理每个列表及其文档
    processed_lists = set()  # 用于跟踪已处理的列表
    for list_item in sort_lists_by_level(selected_lists):
        if list_item.id in processed_lists:
            continue

        # 添加列表标题
        level = list_item.level_id + 1  # Word文档的标题级别从1开始
        heading = doc.add_heading(f"{'  ' * list_item.level_id}{list_item.name}", level=level)
        heading.style.font.name = '宋体'
        
        # 获取列表下的文档
        documents = db.query(models.Document).filter(
            models.Document.list_id == list_item.id,
            models.Document.is_deleted == False
        ).all()

        # 添加文档内容
        for doc_item in documents:
            # 添加文档标题
            doc_heading = doc.add_heading(f"{'  ' * (list_item.level_id + 1)}{doc_item.title}", level=level + 1)
            doc_heading.style.font.name = '宋体'
            
            # 转换并添加文档内容
            content_paragraph = doc.add_paragraph()
            plain_text = convert_html_to_text(doc_item.content)
            content_paragraph.add_run(plain_text)
            content_paragraph.style.font.name = '宋体'
            
            # 添加文档元数据
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.add_run("文档信息：").bold = True
            meta_paragraph.add_run(" | ")
            meta_paragraph.add_run(f"创建时间：{doc_item.created_at.strftime('%Y-%m-%d %H:%M:%S')} | ")
            meta_paragraph.add_run(f"最后修改：{doc_item.updated_at.strftime('%Y-%m-%d %H:%M:%S')} | ")
            meta_paragraph.add_run(f"当前版本：{doc_item.version}")
            meta_paragraph.style.font.name = '宋体'
            
            # 添加分隔线
            doc.add_paragraph("=" * 50)
            doc.add_paragraph()  # 添加空行

        processed_lists.add(list_item.id)

    # 保存文档
    file_name = f"export_{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join("exports", file_name)
    
    # 确保exports目录存在
    os.makedirs("exports", exist_ok=True)
    
    # 保存文件
    doc.save(file_path)
    
    # 返回文件
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name
    )


###############     文档搜索      ##################

@app.get("/documents/search/", response_model=List[schemas.Document])
def search_documents(
    keyword: Optional[str] = None,
    project_id: Optional[int] = None,
    list_id: Optional[int] = None,
    start_date: Optional[datetime] = None,
    end_date: Optional[datetime] = None,
    skip: int = 0,
    limit: int = 10,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """搜索文档
    
    参数:
    - keyword: 搜索关键词（标题或内容）
    - project_id: 项目ID
    - list_id: 列表ID
    - start_date: 开始日期
    - end_date: 结束日期
    - skip: 分页起始位置
    - limit: 每页数量
    """
    # 构建查询条件
    query = db.query(models.Document).filter(models.Document.is_deleted == False)
    
    # 添加关键词搜索
    if keyword:
        query = query.filter(
            or_(
                models.Document.title.ilike(f"%{keyword}%"),
                models.Document.content.ilike(f"%{keyword}%")
            )
        )
    
    # 添加项目筛选
    if project_id:
        query = query.join(models.mylist).filter(models.mylist.project_id == project_id)
    
    # 添加列表筛选
    if list_id:
        query = query.filter(models.Document.list_id == list_id)
    
    # 添加日期范围筛选
    if start_date:
        query = query.filter(models.Document.created_at >= start_date)
    if end_date:
        query = query.filter(models.Document.created_at <= end_date)
    
    # 添加权限检查
    query = query.join(models.mylist).join(models.Project).filter(
        models.Project.owner_id == current_user.id
    )
    
    # 按创建时间倒序排序
    query = query.order_by(models.Document.created_at.desc())
    
    # 执行分页查询
    total = query.count()
    documents = query.offset(skip).limit(limit).all()
    
    return documents


###############     文档权限管理      ##################

@app.post("/documents/{document_id}/permissions/", response_model=schemas.DocumentPermission)
def create_document_permission(
    document_id: int,
    permission: schemas.DocumentPermissionCreate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """创建文档权限"""
    # 检查当前用户是否有管理权限
    if not crud.check_document_permission(db, document_id, current_user.id, PermissionLevel.ADMIN):
        raise HTTPException(status_code=403, detail="没有权限管理此文档")
    
    # 检查目标用户是否存在
    target_user = crud.get_user(db, user_id=permission.user_id)
    if not target_user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    # 创建权限
    db_permission = crud.create_document_permission(
        db=db,
        document_id=document_id,
        user_id=permission.user_id,
        permission_level=permission.permission_level
    )
    return db_permission

@app.put("/documents/{document_id}/permissions/{user_id}/", response_model=schemas.DocumentPermission)
def update_document_permission(
    document_id: int,
    user_id: int,
    permission: schemas.DocumentPermissionBase,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """更新文档权限"""
    # 检查当前用户是否有管理权限
    if not crud.check_document_permission(db, document_id, current_user.id, PermissionLevel.ADMIN):
        raise HTTPException(status_code=403, detail="没有权限管理此文档")
    
    # 更新权限
    db_permission = crud.update_document_permission(
        db=db,
        document_id=document_id,
        user_id=user_id,
        permission_level=permission.permission_level
    )
    if not db_permission:
        raise HTTPException(status_code=404, detail="权限记录不存在")
    return db_permission

@app.delete("/documents/{document_id}/permissions/{user_id}/")
def delete_document_permission(
    document_id: int,
    user_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """删除文档权限"""
    # 检查当前用户是否有管理权限
    if not crud.check_document_permission(db, document_id, current_user.id, PermissionLevel.ADMIN):
        raise HTTPException(status_code=403, detail="没有权限管理此文档")
    
    # 删除权限
    if not crud.delete_document_permission(db=db, document_id=document_id, user_id=user_id):
        raise HTTPException(status_code=404, detail="权限记录不存在")
    return {"message": "权限删除成功"}

@app.get("/documents/{document_id}/permissions/", response_model=List[schemas.DocumentPermission])
def get_document_permissions(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """获取文档的所有权限"""
    # 检查当前用户是否有管理权限
    if not crud.check_document_permission(db, document_id, current_user.id, PermissionLevel.ADMIN):
        raise HTTPException(status_code=403, detail="没有权限查看此文档的权限设置")
    
    return crud.get_document_permissions(db=db, document_id=document_id)

###############     文档协作      ##################

@app.post("/documents/{document_id}/lock/", response_model=schemas.DocumentLock)
def acquire_lock(
    document_id: int,
    lock_duration: int = 30,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """获取文档编辑锁"""
    # 检查文档权限
    if not crud.check_document_permission(db, document_id, current_user.id, PermissionLevel.WRITE):
        raise HTTPException(status_code=403, detail="没有编辑权限")
    
    # 获取编辑锁
    lock = crud.acquire_document_lock(
        db=db,
        document_id=document_id,
        user_id=current_user.id,
        lock_duration=lock_duration
    )
    
    if not lock:
        raise HTTPException(status_code=409, detail="文档已被其他用户锁定")
    return lock

@app.delete("/documents/{document_id}/lock/")
def release_lock(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(auth.get_current_active_user)
):
    """释放文档编辑锁"""
    if not crud.release_document_lock(db, document_id, current_user.id):
        raise HTTPException(status_code=404, detail="未找到有效的锁定记录")
    return {"message": "锁定已释放"}

###############     预留Item      ##################


@app.get("/")
def read_root():
    return {"Hello": "World"}


@app.get("/items/{item_id}", response_model=schemas.Item)
def read_item(item_id: int, db: Session = Depends(get_db)):
    db_item = crud.get_item(db, item_id=item_id)
    if db_item is None:
        raise HTTPException(status_code=404, detail="Item not found")
    return db_item


@app.get("/items/", response_model=list[schemas.Item])
def read_items(skip: int = 0, limit: int = 10, db: Session = Depends(get_db)):
    items = crud.get_items(db, skip=skip, limit=limit)
    return items


@app.post("/items/", response_model=schemas.Item)
def create_item(item: schemas.ItemCreate, db: Session = Depends(get_db)):
    return crud.create_item(db=db, item=item)

@app.exception_handler(exceptions.DocumentNotFoundError)
async def document_not_found_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail}
    )

@app.exception_handler(exceptions.PermissionDeniedError)
async def permission_denied_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail}
    )

@app.exception_handler(exceptions.DocumentLockedError)
async def document_locked_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail}
    )

@app.exception_handler(exceptions.DatabaseError)
async def database_error_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail}
    )

@app.exception_handler(exceptions.ValidationError)
async def validation_error_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail}
    )

@app.exception_handler(SQLAlchemyError)
async def sqlalchemy_error_handler(request, exc):
    return JSONResponse(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        content={"detail": "数据库操作失败"}
    )
